"""
Text extraction from PDF, DOCX, and XLSX files plus word-level chunking.
All parsers return plain text; the chunker splits it into overlapping windows.
"""
import io
from dataclasses import dataclass, field
from typing import List, Optional

from loguru import logger

CHUNK_SIZE = 500    # words per chunk
CHUNK_OVERLAP = 50  # words shared between adjacent chunks


@dataclass
class ParsedDocument:
    title: str
    text: str
    metadata: dict = field(default_factory=dict)


# ── individual parsers ────────────────────────────────────────────────────────

def parse_pdf(data: bytes, title: str) -> ParsedDocument:
    text = ""
    num_pages = 0

    # Attempt 1: pdfplumber — text + tables (best for datasheets and price lists)
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            num_pages = len(pdf.pages)
            parts = []
            for page in pdf.pages:
                # Plain text (paragraphs, labels, annotations on graphs)
                t = page.extract_text()
                if t and t.strip():
                    parts.append(t.strip())
                # Tables — specs tables, parameter tables, price rows
                for table in page.extract_tables():
                    for row in table:
                        cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                        # Deduplicate merged cells
                        seen: set = set()
                        unique: List[str] = []
                        for c in cells:
                            if c not in seen:
                                seen.add(c)
                                unique.append(c)
                        if unique:
                            parts.append(" | ".join(unique))
            text = "\n\n".join(parts)
        if text.strip():
            logger.info(f"[PARSER] pdfplumber OK: '{title}' — {num_pages} pages, {len(text)} chars")
            return ParsedDocument(title=title, text=text, metadata={"pages": num_pages, "parser": "pdfplumber"})
        logger.warning(f"[PARSER] pdfplumber returned empty text for '{title}' — trying pymupdf")
    except Exception as e:
        logger.warning(f"[PARSER] pdfplumber failed for '{title}': {e} — trying pymupdf")

    # Attempt 2: pymupdf (fitz) — handles more layouts
    try:
        import fitz  # pymupdf
        doc = fitz.open(stream=data, filetype="pdf")
        num_pages = len(doc)
        parts = [doc[i].get_text() for i in range(num_pages)]
        text = "\n\n".join(p.strip() for p in parts if p.strip())
        doc.close()
        if text.strip():
            logger.info(f"[PARSER] pymupdf OK: '{title}' — {num_pages} pages, {len(text)} chars")
            return ParsedDocument(title=title, text=text, metadata={"pages": num_pages, "parser": "pymupdf"})
        logger.warning(f"[PARSER] pymupdf returned empty text for '{title}' — trying PyPDF2")
    except Exception as e:
        logger.warning(f"[PARSER] pymupdf failed for '{title}': {e} — trying PyPDF2")

    # Attempt 3: PyPDF2 — fallback
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(data))
        num_pages = len(reader.pages)
        pages = [p.extract_text() for p in reader.pages if p.extract_text()]
        text = "\n\n".join(p.strip() for p in pages)
        if text.strip():
            logger.info(f"[PARSER] PyPDF2 OK: '{title}' — {num_pages} pages, {len(text)} chars")
            return ParsedDocument(title=title, text=text, metadata={"pages": num_pages, "parser": "PyPDF2"})
    except Exception as e:
        logger.warning(f"[PARSER] PyPDF2 failed for '{title}': {e}")

    logger.error(f"[PARSER] All PDF parsers failed for '{title}' — file may be scanned image (no text layer)")
    return ParsedDocument(title=title, text="", metadata={"pages": num_pages, "parser": "none"})


def parse_docx(data: bytes, title: str) -> ParsedDocument:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts: List[str] = []

    # Paragraphs (headings, plain text)
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    # Tables — price lists and specs are almost always here
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            # deduplicate merged cells (python-docx repeats them)
            seen, unique = set(), []
            for c in cells:
                if c not in seen:
                    seen.add(c)
                    unique.append(c)
            if unique:
                parts.append(" | ".join(unique))

    return ParsedDocument(title=title, text="\n\n".join(parts))


def parse_xlsx(data: bytes, title: str) -> ParsedDocument:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    rows: List[str] = []
    for sheet in wb.worksheets:
        rows.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                rows.append(" | ".join(cells))
    return ParsedDocument(title=title, text="\n".join(rows))


_PARSERS = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "xlsx": parse_xlsx,
}


def parse_document(data: bytes, filename: str) -> Optional[ParsedDocument]:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    parser = _PARSERS.get(ext)
    if parser is None:
        return None
    return parser(data, filename)


# ── chunker ───────────────────────────────────────────────────────────────────

def chunk_text(
    text: str,
    size: int = CHUNK_SIZE,
    overlap: int = CHUNK_OVERLAP,
) -> List[str]:
    words = text.split()
    chunks: List[str] = []
    i = 0
    while i < len(words):
        piece = " ".join(words[i : i + size])
        if piece.strip():
            chunks.append(piece)
        i += size - overlap
    return chunks
